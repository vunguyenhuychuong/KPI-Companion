"""Xuat du lieu KPI ra nhieu dinh dang (CSV/Markdown/JSON/Excel/PDF/DOCX) + bao cao gui quan ly.

- Nguoi dung chon DINH DANG (formats) + PHAM VI du lieu (sections).
- 1 file -> tai thang; nhieu file -> dong goi .zip.
- Tieng Viet trong PDF: dung TTF Unicode cua Windows (Arial/Segoe UI/Tahoma) — co fallback.
"""
import csv
import io
import json
import os
import zipfile
from datetime import date

from sqlalchemy import select
from sqlalchemy.orm import Session

from .. import models, schemas
from . import kpi_service

VALID_FORMATS = ["csv", "md", "json", "xlsx", "pdf", "docx"]
VALID_SECTIONS = ["kpis", "work_items", "changelog", "reports"]

SECTION_TITLES = {
    "kpis": "KPI & Mục tiêu",
    "work_items": "Đầu việc",
    "changelog": "Lịch sử thay đổi",
    "reports": "Báo cáo kỳ đã lưu",
}

_HEALTH_VN = {"green": "Đúng tiến độ", "yellow": "Cần chú ý", "red": "Rủi ro"}

# Font Unicode (co dau tieng Viet) cho PDF — thu lan luot, lay cai dau tien co tren may
_PDF_FONT_CANDIDATES = [
    (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
    (r"C:\Windows\Fonts\segoeui.ttf", r"C:\Windows\Fonts\segoeuib.ttf"),
    (r"C:\Windows\Fonts\tahoma.ttf", r"C:\Windows\Fonts\tahomabd.ttf"),
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
]


# ---------- Thu thap du lieu theo section ----------
def _kpi_rows(db: Session, user_id: int) -> list[dict]:
    today = date.today()
    rows = []
    for k in kpi_service.get_active_kpis(db, user_id):
        health, gap = kpi_service.health_of(k, today)
        exp = kpi_service.expected_progress(k, today)
        rows.append({
            "Mục tiêu": k.objective_name or "(chưa gắn)",
            "KPI": k.name,
            "Phân loại": "Cá nhân" if k.category == "Personal" else "Công việc",
            "Đơn vị": k.unit,
            "Chỉ tiêu": k.target_value,
            "Thực đạt": k.current_value,
            "Tiến độ (%)": k.progress,
            "Kỳ vọng (%)": exp,
            "Chênh lệch": gap,
            "Trạng thái": _HEALTH_VN.get(health, health),
            "Trọng số (%)": k.weight,
            "Deadline": str(k.deadline or f"{k.year}-12-31"),
        })
    return rows


def _work_item_rows(db: Session, user_id: int) -> list[dict]:
    items = list(db.scalars(
        select(models.WorkItem)
        .where(models.WorkItem.user_id == user_id, models.WorkItem.confirmed == True)  # noqa: E712
        .order_by(models.WorkItem.work_date.desc().nullslast(), models.WorkItem.created_at.desc())
    ))
    return [{
        "Ngày": str(w.work_date or w.created_at.date()),
        "Đầu việc": w.title,
        "Trạng thái": schemas.STATUS_LABELS.get(w.status, w.status),
        "KPI": w.kpi.name if w.kpi else "(chưa gắn KPI)",
        "Thay đổi": w.progress_delta,
        "Nguồn": w.source,
        "Nguồn gốc": w.source_ref,
    } for w in items]


def _changelog_rows(db: Session, user_id: int) -> list[dict]:
    logs = list(db.scalars(
        select(models.KPIChangeLog)
        .join(models.KPI)
        .where(models.KPI.user_id == user_id)
        .order_by(models.KPIChangeLog.changed_at.desc())
    ))
    return [{
        "Thời gian": lg.changed_at.isoformat(sep=" ")[:19],
        "KPI": lg.kpi_name or "KPI đã bị xóa hoặc không còn truy cập được",
        "Trường": lg.field,
        "Giá trị cũ": lg.old_value,
        "Giá trị mới": lg.new_value,
        "Lý do": lg.reason,
    } for lg in logs]


def _report_rows(db: Session, user_id: int) -> list[dict]:
    reports = list(db.scalars(
        select(models.SavedReport)
        .where(models.SavedReport.user_id == user_id)
        .order_by(models.SavedReport.created_at.desc())
    ))
    return [{
        "Kỳ": r.period_label,
        "Ngày tạo": r.created_at.isoformat(sep=" ")[:19],
        "Nội dung": r.content,
    } for r in reports]


_SECTION_BUILDERS = {
    "kpis": _kpi_rows,
    "work_items": _work_item_rows,
    "changelog": _changelog_rows,
    "reports": _report_rows,
}


def _gather(db: Session, user_id: int, sections: list[str]) -> dict[str, list[dict]]:
    return {s: _SECTION_BUILDERS[s](db, user_id) for s in sections if s in _SECTION_BUILDERS}


# ---------- Renderers ----------
def _csv_bytes(rows: list[dict]) -> bytes:
    buf = io.StringIO()
    if rows:
        writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()), extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)
    else:
        buf.write("(không có dữ liệu)\n")
    return buf.getvalue().encode("utf-8-sig")  # BOM de Excel mo dung tieng Viet


def _md_table(rows: list[dict]) -> str:
    if not rows:
        return "_(không có dữ liệu)_\n"
    cols = list(rows[0].keys())
    out = ["| " + " | ".join(cols) + " |", "| " + " | ".join("---" for _ in cols) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(r.get(c, "")).replace("\n", " ").replace("|", "\\|") for c in cols) + " |")
    return "\n".join(out) + "\n"


def _render_md(data: dict[str, list[dict]]) -> bytes:
    parts = [f"# Xuất dữ liệu KPI Companion\n\n_Ngày xuất: {date.today().isoformat()}_\n"]
    for sec, rows in data.items():
        parts.append(f"\n## {SECTION_TITLES.get(sec, sec)}\n")
        if sec == "reports":
            for r in rows:
                parts.append(f"\n### {r['Kỳ']} ({r['Ngày tạo']})\n\n{r['Nội dung']}\n")
            if not rows:
                parts.append("_(không có dữ liệu)_\n")
        else:
            parts.append(_md_table(rows))
    return "\n".join(parts).encode("utf-8")


def _render_json(data: dict[str, list[dict]]) -> bytes:
    payload = {"exported_at": date.today().isoformat(), "sections": data}
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def _render_xlsx(data: dict[str, list[dict]]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    wb.remove(wb.active)

    title_fill = PatternFill("solid", fgColor="1F4E79")
    title_font = Font(bold=True, color="FFFFFF", size=14)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    header_font = Font(bold=True, color="17365D")
    thin = Side(style="thin", color="D7DEE8")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    if not data:
        data = {"kpis": []}

    for sec, rows in data.items():
        title = SECTION_TITLES.get(sec, sec)
        ws = wb.create_sheet(title[:31])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
        ws["A1"] = title
        ws["A1"].fill = title_fill
        ws["A1"].font = title_font
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 26
        ws["A2"] = f"Ngày xuất: {date.today().isoformat()}"
        ws["A2"].font = Font(italic=True, color="666666")

        cols = list(rows[0].keys()) if rows else ["Thông tin"]
        for col_idx, col_name in enumerate(cols, 1):
            cell = ws.cell(row=4, column=col_idx, value=col_name)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        if rows:
            for row in rows:
                ws.append([row.get(col, "") for col in cols])
        else:
            ws.append(["Không có dữ liệu"])

        for row in ws.iter_rows(min_row=5, max_row=ws.max_row, max_col=len(cols)):
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(vertical="top", wrap_text=True)

        ws.freeze_panes = "A5"
        ws.auto_filter.ref = ws.dimensions
        for idx, col in enumerate(cols, 1):
            values = [str(col), *[str(r.get(col, "")) for r in rows[:80]]]
            width = min(max(len(v) for v in values) + 3, 48)
            ws.column_dimensions[ws.cell(row=4, column=idx).column_letter].width = max(width, 12)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _find_pdf_font() -> tuple[str | None, str | None]:
    for reg, bold in _PDF_FONT_CANDIDATES:
        if os.path.exists(reg):
            return reg, (bold if os.path.exists(bold) else None)
    return None, None


def _render_pdf(data: dict[str, list[dict]]) -> bytes:
    from fpdf import FPDF

    reg, bold = _find_pdf_font()
    pdf = FPDF()
    pdf.set_auto_page_break(True, margin=15)
    if reg:
        pdf.add_font("Main", "", reg)
        pdf.add_font("Main", "B", bold or reg)
        fam = "Main"
    else:
        fam = "Helvetica"  # fallback: khong co font Unicode -> tieng Viet co the mat dau
    pdf.add_page()
    pdf.set_font(fam, "B", 16)
    pdf.cell(0, 10, "Xuất dữ liệu KPI Companion", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(fam, "", 10)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 7, f"Ngày xuất: {date.today().isoformat()}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    epw = pdf.w - 2 * pdf.l_margin
    for sec, rows in data.items():
        pdf.ln(3)
        pdf.set_font(fam, "B", 13)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(0, 9, SECTION_TITLES.get(sec, sec), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font(fam, "", 10)
        if not rows:
            pdf.cell(0, 6, "(không có dữ liệu)", new_x="LMARGIN", new_y="NEXT")
            continue
        if sec == "reports":
            for r in rows:
                pdf.set_font(fam, "B", 11)
                pdf.multi_cell(epw, 6, f"{r['Kỳ']} ({r['Ngày tạo']})")
                pdf.set_font(fam, "", 10)
                pdf.multi_cell(epw, 5, str(r["Nội dung"]))
                pdf.ln(2)
        else:
            for r in rows:
                line = "  •  ".join(f"{k}: {v}" for k, v in r.items())
                pdf.multi_cell(epw, 5, line, border="B")
    out = pdf.output()
    return bytes(out)


def _render_docx(data: dict[str, list[dict]]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Xuất dữ liệu KPI Companion", level=0)
    doc.add_paragraph(f"Ngày xuất: {date.today().isoformat()}")
    for sec, rows in data.items():
        doc.add_heading(SECTION_TITLES.get(sec, sec), level=1)
        if not rows:
            doc.add_paragraph("(không có dữ liệu)")
            continue
        if sec == "reports":
            for r in rows:
                doc.add_heading(f"{r['Kỳ']} ({r['Ngày tạo']})", level=2)
                doc.add_paragraph(str(r["Nội dung"]))
        else:
            cols = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for r in rows:
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(r.get(c, ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- Diem vao chinh ----------
def build_export(
    db: Session, user_id: int, sections: list[str], formats: list[str]
) -> tuple[str, bytes, str]:
    """Tra ve (filename, content, media_type). 1 file -> file goc; >1 -> .zip."""
    sections = [s for s in sections if s in VALID_SECTIONS] or ["kpis"]
    formats = [f for f in formats if f in VALID_FORMATS] or ["csv"]
    data = _gather(db, user_id, sections)
    stamp = date.today().isoformat()

    files: list[tuple[str, bytes]] = []  # (ten file, noi dung)
    for fmt in formats:
        if fmt == "csv":
            for sec, rows in data.items():
                files.append((f"kpi-{sec}-{stamp}.csv", _csv_bytes(rows)))
        elif fmt == "md":
            files.append((f"kpi-export-{stamp}.md", _render_md(data)))
        elif fmt == "json":
            files.append((f"kpi-export-{stamp}.json", _render_json(data)))
        elif fmt == "xlsx":
            suffix = "-".join(sections) if len(sections) <= 2 else "data"
            files.append((f"kpi-{suffix}-{stamp}.xlsx", _render_xlsx(data)))
        elif fmt == "pdf":
            files.append((f"kpi-export-{stamp}.pdf", _render_pdf(data)))
        elif fmt == "docx":
            files.append((f"kpi-export-{stamp}.docx", _render_docx(data)))

    if len(files) == 1:
        name, content = files[0]
        return name, content, _MEDIA.get(name.rsplit(".", 1)[-1], "application/octet-stream")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in files:
            zf.writestr(name, content)
    return f"kpi-export-{stamp}.zip", buf.getvalue(), "application/zip"


_MEDIA = {
    "csv": "text/csv; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "json": "application/json",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "zip": "application/zip",
}


# ---------- Accountability Proxy: bao cao gui quan ly (loc rieng tu) ----------
def build_manager_report(db: Session, user_id: int) -> str:
    """Bao cao tom tat dang markdown CHI gom KPI Cong viec (an KPI Ca nhan — co lap ngu canh M5)."""
    today = date.today()
    kpis = [k for k in kpi_service.get_active_kpis(db, user_id) if k.category != "Personal"]
    lines = [
        f"# Báo cáo tiến độ KPI — {today.isoformat()}",
        "",
        "_(Chỉ gồm KPI công việc; các mục tiêu cá nhân được giữ riêng tư.)_",
        "",
        "## Tổng quan KPI công việc",
    ]
    if not kpis:
        lines.append("\n(Chưa có KPI công việc nào.)")
        return "\n".join(lines)
    for k in kpis:
        health, gap = kpi_service.health_of(k, today)
        exp = kpi_service.expected_progress(k, today)
        lines.append(
            f"- **{k.name}** — {k.current_value:g}/{k.target_value:g} {k.unit} "
            f"= {k.progress:.0f}% (kỳ vọng {exp:.0f}%, lệch {gap:+.0f}%, {_HEALTH_VN.get(health, health)})"
        )
    behind = [k for k in kpis if kpi_service.health_of(k, today)[0] != "green"]
    if behind:
        lines.append("\n## ⚠️ Cần chú ý")
        for k in behind:
            _, gap = kpi_service.health_of(k, today)
            lines.append(f"- **{k.name}**: đang chậm {abs(gap):.0f}% so với kế hoạch.")
    return "\n".join(lines)
